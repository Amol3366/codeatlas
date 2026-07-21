"""Reader coverage for code, docs, notebooks, DOCX, and text fallback."""

from __future__ import annotations

import json
from pathlib import Path

from app.ingestion.classify import classify
from app.ingestion.readers import read_source_file


def test_classify_common_code_and_document_extensions() -> None:
    assert classify(Path("main.py")) == ("code", "python")
    assert classify(Path("app.tsx")) == ("code", "tsx")
    assert classify(Path("notebook.ipynb")) == ("doc", None)
    assert classify(Path("manual.pdf")) == ("doc", None)
    assert classify(Path("spec.docx")) == ("doc", None)
    assert classify(Path("custom.projectnote")) == ("doc", None)
    assert classify(Path("image.png")) is None


def test_read_source_file_reads_unknown_text_extension(tmp_path: Path) -> None:
    path = tmp_path / "notes.projectnote"
    path.write_text("Project-specific text format", encoding="utf-8")

    source = read_source_file(path, tmp_path, "repo", None)

    assert source is not None
    assert source.kind == "doc"
    assert source.path == "notes.projectnote"
    assert "Project-specific" in source.content


def test_read_source_file_skips_binary_unknown_extension(tmp_path: Path) -> None:
    path = tmp_path / "payload.custom"
    path.write_bytes(b"\x00\x01\x02\x03")

    assert read_source_file(path, tmp_path, "repo", None) is None


def test_read_source_file_flattens_notebook(tmp_path: Path) -> None:
    path = tmp_path / "analysis.ipynb"
    path.write_text(
        json.dumps(
            {
                "cells": [
                    {"cell_type": "markdown", "source": ["# Title\n", "Notes"]},
                    {"cell_type": "code", "source": ["print('hello')"]},
                ]
            }
        ),
        encoding="utf-8",
    )

    source = read_source_file(path, tmp_path, "repo", None)

    assert source is not None
    assert source.kind == "doc"
    assert "# Title" in source.content
    assert "print('hello')" in source.content


def test_read_source_file_extracts_docx(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "requirements.docx"
    document = Document()
    document.add_paragraph("Project requirements")
    document.add_paragraph("Index source code and documentation.")
    document.save(path)

    source = read_source_file(path, tmp_path, "repo", None)

    assert source is not None
    assert source.kind == "doc"
    assert source.path == "requirements.docx"
    assert "Project requirements" in source.content
    assert "Index source code" in source.content

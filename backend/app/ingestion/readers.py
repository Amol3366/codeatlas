"""Reading file contents into :class:`SourceFile` objects (CLAUDE.md §4).

Text source and document files are read as UTF-8; Jupyter notebooks are
flattened to their markdown + code cells. PDF and DOCX files are parsed with
dedicated readers.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path

from app.ingestion.classify import classify
from app.models import SourceFile

logger = logging.getLogger(__name__)

_BINARY_SNIFF_BYTES = 4096


def _read_notebook(path: Path) -> str | None:
    """Flatten a Jupyter notebook into markdown + code cell text."""
    try:
        notebook = json.loads(path.read_text(encoding="utf-8", errors="strict"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        logger.warning("Skipping unreadable notebook %s: %s", path, exc)
        return None
    parts: list[str] = []
    for cell in notebook.get("cells", []):
        source = cell.get("source", "")
        text = "".join(source) if isinstance(source, list) else str(source)
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _read_pdf(path: Path) -> str | None:
    """Extract text from a PDF, one page per section."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"# Page {index}\n\n{text.strip()}")
        return "\n\n".join(parts)
    except Exception as exc:  # noqa: BLE001 - malformed PDFs should not kill ingestion
        logger.warning("Skipping unreadable PDF %s: %s", path, exc)
        return None


def _read_docx(path: Path) -> str | None:
    """Extract text from DOCX paragraphs and tables."""
    try:
        from docx import Document

        document = Document(str(path))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception as exc:  # noqa: BLE001 - malformed DOCX should not kill ingestion
        logger.warning("Skipping unreadable DOCX %s: %s", path, exc)
        return None


def _looks_binary(path: Path) -> bool:
    """Return True for files that appear binary from an initial byte sample."""
    try:
        with path.open("rb") as handle:
            sample = handle.read(_BINARY_SNIFF_BYTES)
    except OSError:
        return True
    if b"\x00" in sample:
        return True
    if not sample:
        return False
    control_bytes = sum(byte < 32 and byte not in (9, 10, 12, 13) for byte in sample)
    return control_bytes / len(sample) > 0.30


def _read_text(path: Path) -> str | None:
    """Read a UTF-8 text file, skipping binary/unreadable files."""
    if _looks_binary(path):
        logger.warning("Skipping binary-looking file %s", path)
        return None
    try:
        return path.read_text(encoding="utf-8", errors="strict")
    except (UnicodeDecodeError, OSError) as exc:
        logger.warning("Skipping unreadable text file %s: %s", path, exc)
        return None


def read_source_file(
    abs_path: Path,
    root: Path,
    repo: str,
    commit_hash: str | None,
) -> SourceFile | None:
    """Read one file into a :class:`SourceFile`, or ``None`` to skip it."""
    classification = classify(abs_path)
    if classification is None:
        return None
    kind, language = classification
    rel_path = abs_path.relative_to(root).as_posix()
    suffix = abs_path.suffix.lower()

    if suffix == ".ipynb":
        content = _read_notebook(abs_path)
        if content is None:
            return None
    elif suffix == ".pdf":
        content = _read_pdf(abs_path)
        if content is None:
            return None
    elif suffix == ".docx":
        content = _read_docx(abs_path)
        if content is None:
            return None
    else:
        content = _read_text(abs_path)
        if content is None:
            return None

    if not content.strip():
        return None

    return SourceFile(
        repo=repo,
        path=rel_path,
        abs_path=str(abs_path),
        content=content,
        language=language,
        kind=kind,
        commit_hash=commit_hash,
    )
